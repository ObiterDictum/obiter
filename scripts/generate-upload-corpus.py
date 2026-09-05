"""Generate the P2.24 upload-acceptance corpus and measure per-entry compression ratios.

Real-toolchain provenance: every fixture is produced by python-docx (genuine
Word-compatible OOXML), with footnotes/numbering/tracked-changes added by
surgery on the underlying lxml tree using constructs Word itself writes.
Fixtures are synthetic and fictional — no client matter data.

Usage:
  python3 scripts/generate-upload-corpus.py --out <dir>   # write fixtures
  python3 scripts/generate-upload-corpus.py --measure <dir>  # report ratios only
"""
from __future__ import annotations

import argparse
import io
import sys
import zipfile

OUT_FILES = [
    "letter-plain.docx",
    "letter-table.docx",
    "letter-footnotes-numbering.docx",
    "letter-tracked-changes.docx",
    "letter-image.docx",
]

LOREM = (
    "Dear Mr Cartwright, We write further to your letter of 14 March regarding "
    "the boundary dispute at Mill Farm. Our client instructs us that the fence "
    "line has stood unaltered since September 2011, and the Land Registry plan "
    "for title number AB123456 supports that position. "
)


def _para(doc, text, style=None):
    from docx import Document  # noqa: F401  (import here for clearer errors)

    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p


def make_plain():
    from docx import Document

    doc = Document()
    _para(doc, "Re: Boundary dispute — Mill Farm", style="Title")
    for i in range(12):
        _para(doc, f"Paragraph {i + 1}. {LOREM * 3}")
    return doc


def make_table():
    from docx import Document

    doc = make_plain()
    doc.add_paragraph("Schedule of correspondence")
    table = doc.add_table(rows=9, cols=4)
    table.style = "Light Shading"
    headers = ("Date", "From", "To", "Subject")
    for j, h in enumerate(headers):
        table.cell(0, j).text = h
    for i in range(1, 9):
        row = (f"2024-0{i}-1{i}", "Seller", "Buyer", f"Replies and enclosures {i} " + LOREM)
        for j, v in enumerate(row):
            table.cell(i, j).text = v
    for i in range(4):
        _para(doc, f"After-table paragraph {i + 1}. {LOREM * 2}")
    return doc


def _save_and_surgery_doc(tmp_path, final_path, surgery):
    """Copy tmp to final, then apply raw-zip surgery Word itself would write.

    python-docx has no footnote/numbering API and its OPC layer rejects
hand-built parts, so parts Word owns (footnotes.xml, numbering.xml) are
injected post-save with matching rels and content types. The document.xml
edits use the same element shapes Word serialises.
    """
    import shutil
    import tempfile
    import zipfile as _zf
    from lxml import etree

    # Rebuild rather than append: appending would leave duplicate entries and
    # readers disagree about which copy wins.
    tmp_out = tempfile.mktemp(suffix=".docx")
    with _zf.ZipFile(tmp_path, "r") as zin:
        replacements: dict[str, bytes] = {}

        class _Facade:
            def read(self, name):
                return zin.read(name)

            def writestr(self, name, data, compress_type=None):
                replacements[name] = data

        surgery(_Facade(), etree)
        with _zf.ZipFile(tmp_out, "w", _zf.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = replacements.pop(item.filename, zin.read(item.filename))
                zout.writestr(item.filename, data, compress_type=_zf.ZIP_DEFLATED)
            for name, data in replacements.items():
                zout.writestr(name, data, compress_type=_zf.ZIP_DEFLATED)
    shutil.move(tmp_out, final_path)


NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "rel": "http://schemas.openxmlformats.org/package/2006/relationships",
    "ct": "http://schemas.openxmlformats.org/package/2006/content-types",
}


def _sub(parent, tag, attrib=None, text=None):
    from lxml import etree

    el = etree.SubElement(parent, tag, attrib or {})
    if text is not None:
        el.text = text
    return el


def _footnotes_surgery(num_paras=10):
    def apply(zf, etree):
        doc_xml = etree.fromstring(zf.read("word/document.xml"))
        # Numbering: tag body paragraphs carrying MARKER-NUM-i.
        for i in range(num_paras):
            for p in doc_xml.findall(".//w:p", NS):
                t = "".join(t.text or "" for t in p.findall(".//w:t", NS))
                if f"MARKER-NUM-{i}" in t:
                    pPr = p.find("w:pPr", NS)
                    if pPr is None:
                        pPr = etree.Element(f"{{{NS['w']}}}pPr")
                        p.insert(0, pPr)
                    numPr = _sub(pPr, f"{{{NS['w']}}}numPr")
                    _sub(numPr, f"{{{NS['w']}}}ilvl", {f"{{{NS['w']}}}val": "0"})
                    _sub(numPr, f"{{{NS['w']}}}numId", {f"{{{NS['w']}}}val": "1"})
        # Footnote references: runs carrying MARKER-FN-2 / MARKER-FN-3.
        for fid in ("2", "3"):
            for p in doc_xml.findall(".//w:p", NS):
                t = "".join(t.text or "" for t in p.findall(".//w:t", NS))
                if f"MARKER-FN-{fid}" in t:
                    for r in p.findall("w:r", NS):
                        _sub(r, f"{{{NS['w']}}}footnoteReference", {f"{{{NS['w']}}}id": fid})
                        break
        zf.writestr("word/document.xml", etree.tostring(doc_xml, xml_declaration=True, encoding="UTF-8", standalone=True), compress_type=zipfile.ZIP_DEFLATED)
        footnotes = etree.Element(f"{{{NS['w']}}}footnotes", nsmap={"w": NS["w"]})
        for fid, kind in (("0", "separator"), ("1", "continuationSeparator")):
            _sub(footnotes, f"{{{NS['w']}}}footnote", {f"{{{NS['w']}}}type": kind, f"{{{NS['w']}}}id": fid})
        notes = {
            "2": "See Three Rivers DC v Bank of England [2004] UKHL 48 on privilege. " + LOREM * 2,
            "3": "Compare the disclosure timetable agreed on 2 May. " + LOREM * 2,
        }
        for fid, text in notes.items():
            fn = _sub(footnotes, f"{{{NS['w']}}}footnote", {f"{{{NS['w']}}}id": fid})
            p = _sub(fn, f"{{{NS['w']}}}p")
            r = _sub(p, f"{{{NS['w']}}}r")
            _sub(r, f"{{{NS['w']}}}t", text=text)
        zf.writestr("word/footnotes.xml", etree.tostring(footnotes, xml_declaration=True, encoding="UTF-8", standalone=True), compress_type=zipfile.ZIP_DEFLATED)
        numbering = etree.Element(f"{{{NS['w']}}}numbering", nsmap={"w": NS["w"]})
        abstract = _sub(numbering, f"{{{NS['w']}}}abstractNum", {f"{{{NS['w']}}}abstractNumId": "0"})
        lvl = _sub(abstract, f"{{{NS['w']}}}lvl", {f"{{{NS['w']}}}ilvl": "0"})
        _sub(lvl, f"{{{NS['w']}}}start", {f"{{{NS['w']}}}val": "1"})
        _sub(lvl, f"{{{NS['w']}}}numFmt", {f"{{{NS['w']}}}val": "decimal"})
        _sub(lvl, f"{{{NS['w']}}}lvlText", {f"{{{NS['w']}}}val": "%1."})
        num = _sub(numbering, f"{{{NS['w']}}}num", {f"{{{NS['w']}}}numId": "1"})
        _sub(num, f"{{{NS['w']}}}abstractNumId", {f"{{{NS['w']}}}val": "0"})
        zf.writestr("word/numbering.xml", etree.tostring(numbering, xml_declaration=True, encoding="UTF-8", standalone=True), compress_type=zipfile.ZIP_DEFLATED)
        rels = etree.fromstring(zf.read("word/_rels/document.xml.rels"))
        _sub(rels, f"{{{NS['rel']}}}Relationship", {
            "Id": "rIdFn1", "Type": f"{NS['r']}/footnotes", "Target": "footnotes.xml",
        })
        _sub(rels, f"{{{NS['rel']}}}Relationship", {
            "Id": "rIdNum1", "Type": f"{NS['r']}/numbering", "Target": "numbering.xml",
        })
        zf.writestr("word/_rels/document.xml.rels", etree.tostring(rels, xml_declaration=True, encoding="UTF-8", standalone=True), compress_type=zipfile.ZIP_DEFLATED)
        types = etree.fromstring(zf.read("[Content_Types].xml"))
        for part, ctype in (
            ("/word/footnotes.xml", "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"),
            ("/word/numbering.xml", "application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml"),
        ):
            _sub(types, f"{{{NS['ct']}}}Override", {"PartName": part, "ContentType": ctype})
        zf.writestr("[Content_Types].xml", etree.tostring(types, xml_declaration=True, encoding="UTF-8", standalone=True), compress_type=zipfile.ZIP_DEFLATED)
    return apply


def make_footnotes_numbering():
    from docx import Document

    doc = Document()
    _para(doc, "Advice on privilege", style="Title")
    for i in range(10):
        _para(doc, f"MARKER-NUM-{i} Numbered advice point {i + 1}. {LOREM * 2}")
        if i % 3 == 0:
            _para(doc, f"MARKER-FN-{'2' if i % 2 == 0 else '3'} Footnote reference {i}.")
    doc._surgery = _footnotes_surgery(num_paras=10)
    return doc


def make_tracked_changes():
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = make_plain()
    # Mark an insertion and a deletion the way Word serialises revision tracking.
    p = doc.paragraphs[2]._p
    ins = OxmlElement("w:ins")
    ins.set(qn("w:author"), "Test Reviewer")
    ins.set(qn("w:date"), "2024-05-01T10:00:00Z")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "Inserted on review: the gatehouse paddock is excluded. " + LOREM
    r.append(t)
    ins.append(r)
    p.append(ins)
    p2 = doc.paragraphs[4]._p
    delete = OxmlElement("w:del")
    delete.set(qn("w:author"), "Test Reviewer")
    delete.set(qn("w:date"), "2024-05-02T10:00:00Z")
    r2 = OxmlElement("w:r")
    t2 = OxmlElement("w:delText")
    t2.text = "Deleted on review: the earlier admission is withdrawn. " + LOREM
    r2.append(t2)
    delete.append(r2)
    p2.append(delete)
    # Revision ids on runs so styles.xml-adjacent parts stay realistic.
    for para in doc.paragraphs[:6]:
        for run in para.runs:
            run._r.set(qn("w:rsidR"), "00A11B2C")
    return doc


def make_image():
    from docx import Document
    from docx.shared import Inches

    try:
        from PIL import Image
    except ImportError:
        sys.exit("Pillow is required: pip install Pillow python-docx")
    img = Image.new("RGB", (640, 480), color=(34, 60, 90))
    pixels = img.load()
    for x in range(0, 640, 8):
        for y in range(0, 480, 8):
            pixels[x, y] = (200, 180 - (x % 64), 120 + (y % 64))
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)

    doc = Document()
    _para(doc, "Site plan exhibit", style="Title")
    _para(doc, f"Background. {LOREM * 4}")
    doc.add_picture(buf, width=Inches(5))
    _para(doc, f"After-image commentary. {LOREM * 4}")
    return doc


BUILDERS = {
    "letter-plain.docx": make_plain,
    "letter-table.docx": make_table,
    "letter-footnotes-numbering.docx": make_footnotes_numbering,
    "letter-tracked-changes.docx": make_tracked_changes,
    "letter-image.docx": make_image,
}


def measure(path):
    """Return per-entry (name, compressed, uncompressed, ratio) rows."""
    rows = []
    with zipfile.ZipFile(path) as zf:
        for info in zf.infolist():
            if info.is_dir():
                continue
            ratio = info.file_size / info.compress_size if info.compress_size else float("inf")
            rows.append((info.filename, info.compress_size, info.file_size, ratio))
    return rows


def report(paths):
    worst = ("", 0, 0, 0.0)
    for path in paths:
        print(f"== {path}")
        for name, comp, uncomp, ratio in measure(path):
            print(f"  {ratio:7.2f}x  {comp:>8} -> {uncomp:>8}  {name}")
            if comp >= 256 and ratio > worst[3]:
                worst = (f"{path}::{name}", comp, uncomp, ratio)
    print(f"\nWORST ratio>=256B-compressed entry: {worst[3]:.2f}x  ({worst[0]} {worst[1]}->{worst[2]})")
    return worst


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--out", help="directory to write fixtures into")
    ap.add_argument("--measure", help="directory of fixtures to measure")
    args = ap.parse_args()
    if args.measure:
        import glob
        import os

        paths = sorted(glob.glob(os.path.join(args.measure, "*.docx")))
        report(paths)
        return
    if not args.out:
        ap.error("pass --out <dir> or --measure <dir>")
    import os

    os.makedirs(args.out, exist_ok=True)
    paths = []
    for name, build in BUILDERS.items():
        path = os.path.join(args.out, name)
        doc = build()
        surgery = getattr(doc, "_surgery", None)
        if surgery is not None:
            tmp = path + ".base.docx"
            doc.save(tmp)
            _save_and_surgery_doc(tmp, path, surgery)
            os.remove(tmp)
        else:
            doc.save(path)
        paths.append(path)
        print(f"wrote {path}")
    report(paths)


if __name__ == "__main__":
    main()
